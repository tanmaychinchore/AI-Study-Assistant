"""
DOCX / DOC document loader using python-docx.

Extracts text paragraph-by-paragraph, detecting headings and preserving
structural information for downstream processing.

Note: python-docx only supports the modern .docx (OpenXML) format.
Legacy .doc files are not natively supported.
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.core.logging import get_logger
from app.loaders.base_loader import BaseLoader
from app.schemas.document import ExtractedPage

logger = get_logger(__name__)


class DOCXLoader(BaseLoader):
    """Extract text from Word DOCX files."""

    supported_extensions = {".docx", ".doc"}

    def extract(self, file_path: Path) -> list[ExtractedPage]:
        """
        Extract text from a DOCX file.

        Paragraphs are grouped into sections by heading.  If the document
        has headings, each heading starts a new ExtractedPage (section).
        If there are no headings, the entire document is returned as a
        single ExtractedPage.

        Parameters
        ----------
        file_path : Path
            Path to the DOCX file.

        Returns
        -------
        list[ExtractedPage]
            One ExtractedPage per section (heading-delimited) or a single
            page for the whole document.
        """
        self.validate_file(file_path)

        ext = file_path.suffix.lower()
        if ext == ".doc":
            raise ValueError(
                f"Legacy .doc format is not supported. "
                f"Please convert '{file_path.name}' to .docx."
            )

        logger.info("DOCX extraction started: %s", file_path.name)

        try:
            doc = DocxDocument(str(file_path))
        except PackageNotFoundError as exc:
            raise RuntimeError(
                f"Failed to open DOCX '{file_path.name}': file may be corrupt. {exc}"
            ) from exc
        except Exception as exc:
            raise RuntimeError(
                f"Failed to open DOCX '{file_path.name}': {exc}"
            ) from exc

        pages: list[ExtractedPage] = []
        current_heading: str | None = None
        current_paragraphs: list[str] = []

        def _flush_section() -> None:
            """Flush accumulated paragraphs into an ExtractedPage."""
            text = "\n".join(current_paragraphs).strip()
            if text:
                pages.append(
                    ExtractedPage(
                        heading=current_heading,
                        text=text,
                        char_count=len(text),
                    )
                )

        has_headings = False

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            is_heading = style_name.startswith("heading")

            if is_heading:
                has_headings = True
                # Flush the previous section
                _flush_section()
                current_heading = para.text.strip() or None
                current_paragraphs = []
            else:
                para_text = para.text.strip()
                if para_text:
                    current_paragraphs.append(para_text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    current_paragraphs.append(" | ".join(row_texts))

        # Flush the last section
        _flush_section()

        # If there were no headings at all, the document is a single page
        if not has_headings and not pages:
            full_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if full_text:
                pages.append(
                    ExtractedPage(
                        text=full_text,
                        char_count=len(full_text),
                    )
                )

        if not pages:
            raise ValueError(
                f"No text could be extracted from DOCX: {file_path.name}"
            )

        logger.info(
            "DOCX extraction complete: %s — %d sections extracted",
            file_path.name,
            len(pages),
        )
        return pages
